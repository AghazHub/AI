"""
extraction.py - Text extraction with spatial and font metadata.

Extracts text from digital PDFs (via PyMuPDF ``page.get_text("dict")``) and
DOCX files (via ``python-docx``), preserving bounding-box coordinates, font
names, sizes, and bold/italic flags.

The public entry-point is :func:`extract_text`, which accepts an
``UploadResult`` and dispatches to the correct backend automatically.

Usage:
    from ResumeParser.resume_ingestion import process_upload
    from ResumeParser.extraction import extract_text

    result = process_upload("path/to/resume.pdf")
    if not result.error and result.has_selectable_text:
        extraction = extract_text(result)
        print(f"Extracted {len(extraction.blocks)} blocks")
        print(extraction.raw_text[:500])
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# -------------- Optional imports (graceful degradation) ---------------

try:
    import fitz  # PyMuPDF
    _HAS_PYMUPDF = True
except ImportError:
    _HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

from ResumeParser.resume_ingestion import FileType, UploadResult


# ══════════════════════════════════════════════════════════════════════
# Types
# ══════════════════════════════════════════════════════════════════════

@dataclass
class TextSpan:
    """
    A single text span — the finest unit of text with consistent formatting.

    In PDFs this corresponds to a ``span`` from PyMuPDF's ``dict`` output.
    In DOCX files a span is created per ``Run``.
    """

    text: str
    """The raw text content."""

    bbox: tuple[float, float, float, float]
    """Bounding box ``(x0, y0, x1, y1)`` in PDF points (1/72 inch)."""

    font_name: Optional[str] = None
    """PostScript name of the typeface (e.g. ``'Helvetica'``)."""

    font_size: Optional[float] = None
    """Font size in points."""

    is_bold: bool = False
    """``True`` when the font flags or run properties indicate bold."""

    is_italic: bool = False
    """``True`` when the font flags indicate italic."""

    color: Optional[int] = None
    """Font color as an integer (RGB hex) — PDF only."""


@dataclass
class TextBlock:
    """
    A cohesive block of text — usually a paragraph, heading, or table cell.

    Every block carries at least one :class:`TextSpan`.
    """

    text: str
    """Full concatenated text of the block."""

    bbox: tuple[float, float, float, float]
    """Bounding box encompassing **all** spans in this block."""

    page_num: int
    """Zero-based page number (0 for DOCX files)."""

    spans: list[TextSpan] = field(default_factory=list)
    """The individual spans that make up this block."""

    block_type: str = "text"
    """
    One of:
    - ``'text'``   – regular paragraph / text block
    - ``'image'``  – embedded image (no extractable text)
    - ``'table'``  – content extracted from a table cell
    - ``'heading'`` – detected heading (DOCX style or large bold PDF text)
    """

    block_id: Optional[int] = None
    """Index of this block on its page (for debugging / ordering)."""


@dataclass
class ExtractionResult:
    """
    Result produced by :func:`extract_text`.
    """

    blocks: list[TextBlock]
    """All extracted blocks in reading order."""

    raw_text: str
    """Plain concatenation of ``block.text`` — convenient for LLM input."""

    page_count: int
    """Number of pages in the source document."""

    error: Optional[str] = None
    """Error message if extraction failed."""


# ══════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════

def extract_text(upload: UploadResult) -> ExtractionResult:
    """
    Extract text blocks from a resume, dispatching to the correct backend
    based on ``upload.file_type``.

    Args:
        upload: An :class:`~ResumeParser.resume_ingestion.UploadResult`
            produced by :func:`~ResumeParser.resume_ingestion.process_upload`.

    Returns:
        An :class:`ExtractionResult` with structured blocks and raw text.

    Raises:
        RuntimeError: When the file type has no extraction backend yet
            (e.g. scans / images).
    """
    if upload.error:
        return ExtractionResult(blocks=[], raw_text="", page_count=0,
                                error=upload.error)

    if upload.file_type in (FileType.PDF_DIGITAL,):
        if not _HAS_PYMUPDF:
            return ExtractionResult(
                blocks=[], raw_text="", page_count=0,
                error="PyMuPDF is required to extract text from PDFs. "
                      "Install it with: pip install PyMuPDF",
            )
        return _extract_pdf(upload.file_path)

    if upload.file_type in (FileType.DOCX,):
        if not _HAS_DOCX:
            return ExtractionResult(
                blocks=[], raw_text="", page_count=0,
                error="python-docx is required to extract text from DOCX files. "
                      "Install it with: pip install python-docx",
            )
        return _extract_docx(upload.file_path)

    # Scanned PDFs and images — not yet supported here (see OCR.py).
    return ExtractionResult(
        blocks=[], raw_text="", page_count=0,
        error=f"Cannot extract text from '{upload.file_type.value}' files. "
              f"Run OCR first.",
    )


# ══════════════════════════════════════════════════════════════════════
# PDF extraction — PyMuPDF
# ══════════════════════════════════════════════════════════════════════

def _extract_pdf(file_path: Path) -> ExtractionResult:
    """Extract text blocks from a digital PDF via ``page.get_text("dict")``."""
    doc = fitz.open(str(file_path))
    blocks: list[TextBlock] = []
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc.load_page(page_num)
        page_dict = page.get_text("dict")

        for raw_block in page_dict.get("blocks", []):
            block_type = raw_block.get("type", 0)  # 0=text, 1=image

            if block_type == 1:  # Image block
                bbox = tuple(raw_block.get("bbox", (0, 0, 0, 0)))
                block = TextBlock(
                    text="[IMAGE]",
                    bbox=bbox,
                    page_num=page_num,
                    spans=[],
                    block_type="image",
                    block_id=raw_block.get("number"),
                )
                blocks.append(block)
                continue

            # ── Text block ────────────────────────────────────────
            lines = raw_block.get("lines", [])
            if not lines:
                continue

            block_spans: list[TextSpan] = []
            block_text_parts: list[str] = []
            # Start with the block's own bounding box, tighten per span
            block_bbox = raw_block.get("bbox", (0, 0, 0, 0))

            for line in lines:
                for span_data in line.get("spans", []):
                    text = span_data.get("text", "")
                    if not text.strip():
                        continue

                    bbox = tuple(span_data.get("bbox", (0, 0, 0, 0)))
                    flags = span_data.get("flags", 0)

                    span = TextSpan(
                        text=text,
                        bbox=bbox,
                        font_name=span_data.get("font"),
                        font_size=span_data.get("size"),
                        is_bold=bool(flags & 2**4),
                        is_italic=bool(flags & 2**1),
                        color=span_data.get("color"),
                    )
                    block_spans.append(span)
                    block_text_parts.append(text)

            if not block_spans:
                continue

            full_text = " ".join(block_text_parts)
            block = TextBlock(
                text=full_text,
                bbox=tuple(block_bbox),
                page_num=page_num,
                spans=block_spans,
                block_type="text",
                block_id=raw_block.get("number"),
            )
            blocks.append(block)

    doc.close()

    raw_text = "\n".join(b.text for b in blocks if b.block_type != "image")
    return ExtractionResult(blocks=blocks, raw_text=raw_text,
                            page_count=page_count)


# ══════════════════════════════════════════════════════════════════════
# DOCX extraction — python-docx
# ══════════════════════════════════════════════════════════════════════

_HEADING_STYLES = frozenset({"Heading 1", "Heading 2", "Heading 3",
                              "Heading 4", "Title", "Subtitle"})


def _extract_docx(file_path: Path) -> ExtractionResult:
    """Extract text blocks from a DOCX file via python-docx."""
    doc = DocxDocument(str(file_path))
    blocks: list[TextBlock] = []

    # ── Paragraphs ────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        is_heading = style_name in _HEADING_STYLES

        spans: list[TextSpan] = []
        for run in para.runs:
            run_text = run.text.strip()
            if not run_text:
                continue

            # python-docx stores font size in EMU; convert to points
            size_pt: Optional[float] = None
            if run.font.size is not None:
                size_pt = run.font.size.pt

            spans.append(TextSpan(
                text=run_text,
                bbox=(0, 0, 0, 0),  # DOCX has no native bounding boxes
                font_name=run.font.name,
                font_size=size_pt,
                is_bold=bool(run.bold),
                is_italic=bool(run.italic),
            ))

        block_type = "heading" if is_heading else "text"
        block = TextBlock(
            text=text,
            bbox=(0, 0, 0, 0),
            page_num=0,
            spans=spans or [TextSpan(text=text, bbox=(0, 0, 0, 0))],
            block_type=block_type,
        )
        blocks.append(block)

    # ── Tables ────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                block = TextBlock(
                    text=cell_text,
                    bbox=(0, 0, 0, 0),
                    page_num=0,
                    spans=[TextSpan(text=cell_text, bbox=(0, 0, 0, 0))],
                    block_type="table",
                )
                blocks.append(block)

    raw_text = "\n".join(b.text for b in blocks)
    # DOCX doesn't have pages; count page-breaks approximately
    page_count = 1
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.xml.count('w:br w:type="page"') > 0:
                page_count += 1

    return ExtractionResult(blocks=blocks, raw_text=raw_text,
                            page_count=page_count)


# ══════════════════════════════════════════════════════════════════════
# CLI entry point (quick manual test)
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    from ResumeParser.resume_ingestion import process_upload

    if len(sys.argv) < 2:
        print("Usage:  python extraction.py <path/to/resume.pdf>")
        sys.exit(1)

    upload = process_upload(sys.argv[1])
    if upload.error:
        print(f"Upload error: {upload.error}")
        sys.exit(1)

    result = extract_text(upload)
    if result.error:
        print(f"Extraction error: {result.error}")
        sys.exit(1)

    print(f"Pages: {result.page_count}")
    print(f"Blocks: {len(result.blocks)}")
    print(f"Characters: {len(result.raw_text)}")
    print("─" * 50)

    for i, block in enumerate(result.blocks[:20]):
        preview = block.text[:80] + "…" if len(block.text) > 80 else block.text
        meta = []
        if block.spans and block.spans[0].font_name:
            meta.append(block.spans[0].font_name)
        if block.spans and block.spans[0].font_size:
            meta.append(f"{block.spans[0].font_size:.1f}pt")
        if block.spans and block.spans[0].is_bold:
            meta.append("bold")
        meta_str = f"  [{', '.join(meta)}]" if meta else ""
        print(f"  [{block.block_type:8s}] p{block.page_num} "
              f"{preview}{meta_str}")

    if len(result.blocks) > 20:
        print(f"  … and {len(result.blocks) - 20} more blocks")
